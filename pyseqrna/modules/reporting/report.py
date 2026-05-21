#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
PySeqRNA Run Reporting Module

This module provides comprehensive report generation for PySeqRNA runs.
The report generator is intentionally filesystem-driven, scanning the pipeline's output
directory for checkpoints, tabulations, plots, and text logs to compile an
integrated run book without requiring upstream stage objects to be resident in memory.

Features:
    - Filesystem-driven automatic inspection of PySeqRNA pipeline run output directories
    - Support for multiple output formats: HTML, Markdown, JSON, DOCX, and PDF
    - Dynamic compilation of quality control, alignment, quantification, and differential expression results
    - Embedded visualization plots and interactive HTML tables
    - Robust metadata harvesting from the run checkpoint manager

Configuration:
    The reporting options are configured with:
    - target run output directory path
    - choice of export formats (HTML, Markdown, PDF, DOCX, JSON)

Dependencies:
    - pandas
    - jinja2 (optional, for HTML templating)
    - weasyprint (optional, for PDF compilation)
    - docx (optional, for Word document output)

Classes:
    ReportContext - Data class containing target paths and configuration options.
    ReportGenerator - Generate HTML, Markdown, JSON, DOCX, and PDF reports for a PySeqRNA run.

Exceptions:
    ReportGenerationError - Raised when a report cannot be generated.

Functions:
    generate_report - Convenience wrapper for report generation.

:Created: May 20, 2021
:Updated: April 28, 2026
:Author: Naveen Duhan
:Version: 1.0.0
"""

from __future__ import annotations

import html
import json
import os
import platform
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence

import pandas as pd

from ...__version__ import __version__
from ...utils.checkpoint_manager import CheckpointManager


class ReportGenerationError(Exception):
    """Raised when a report cannot be generated."""


STAGE_NARRATIVES = {
    "quality": (
        "Raw read quality control checks FASTQ files for per-base sequence "
        "quality, GC content, sequence duplication, adapter content, and other "
        "common high-throughput sequencing issues."
    ),
    "trimming": (
        "Read trimming removes low-quality bases and adapter sequence before "
        "alignment. The cleaned reads produced here are normally used by the "
        "alignment stage."
    ),
    "quality_trim": (
        "Post-trimming quality control repeats quality assessment after read "
        "cleaning so the user can confirm that trimming improved or preserved "
        "read quality."
    ),
    "alignment": (
        "Alignment maps reads to the reference genome and stores coordinate "
        "sorted or unsorted BAM files depending on the selected aligner and "
        "tool parameters."
    ),
    "multimapped_groups": (
        "Multimapped-group analysis summarizes reads that map to multiple "
        "features and can be used as an alternate count unit for repetitive or "
        "duplicated genomic regions."
    ),
    "quantification": (
        "Quantification converts alignments into a gene-by-sample count matrix. "
        "This matrix is the main input for normalization and differential "
        "expression analysis."
    ),
    "normalization": (
        "Normalization adjusts count data for library-size, gene-length, or "
        "composition effects depending on the selected method. Normalized "
        "counts are useful for visualization and exploratory analysis."
    ),
    "sample_clustering": (
        "Sample clustering groups samples by expression similarity and can "
        "reveal outliers, batch structure, and condition-specific patterns."
    ),
    "coexpression": (
        "Co-expression analysis groups genes with similar expression profiles "
        "across samples for downstream biological interpretation."
    ),
    "differential": (
        "Differential expression analysis tests pairwise comparisons and "
        "produces all-gene statistics, filtered DEG tables, and DEG gene lists "
        "for downstream interpretation."
    ),
    "visualization": (
        "Visualization summarizes normalized counts and differential expression "
        "results with PCA, t-SNE, heatmaps, MA plots, volcano plots, and related "
        "figures when the required input files are available."
    ),
    "gene_ontology": (
        "Gene Ontology enrichment identifies over-represented biological "
        "processes, molecular functions, and cellular components among DEG lists."
    ),
    "pathway_enrichment": (
        "Pathway enrichment identifies over-represented KEGG pathways among DEG lists for supported Ensembl gene identifiers."
    ),
}

STAGE_COMPLETION_EVIDENCE = {
    "quality": ("fastqc_results", "_fastqc.html", "_fastqc.zip"),
    "trimming": (
        "trim_galore_results",
        "trimming_statistics",
        "_trimming_report.txt",
        "_val_1",
        "_val_2",
    ),
    "quality_trim": ("trim_fastqc_results", "_val_1_fastqc", "_val_2_fastqc"),
    "alignment": (
        "hisat2_results",
        "star_results",
        "alignment_statistics",
        "alignment_stats",
        "_aligned.bam",
    ),
    "multimapped_groups": ("multimapped", "mmg", "Raw_MMG"),
    "quantification": ("Raw_Counts", "raw_counts", "featurecounts", "htseq"),
    "normalization": (
        "normalized_counts",
        "RPKM",
        "FPKM",
        "TPM",
        "CPM",
        "Median_ratio",
        "TMM",
    ),
    "sample_clustering": ("Cluster", "clustering", "dendrogram", "Sample_cluster"),
    "coexpression": ("Coexpression", "Clust", "coexpression", "clusters"),
    "differential": ("All_gene_expression", "Filtered_DEGs", "diff_genes"),
    "visualization": (
        "Volcano_Plots",
        "MA_Plots",
        "Heatmaps",
        "Sample_Plots",
        "Venn_Plots",
        "_volcano",
        "_ma",
        "PCA",
        "t-SNE",
    ),
    "gene_ontology": ("Gene_Ontology", "_GO_", "GO_all", "GO_BP", "GO_MF", "GO_CC"),
    "pathway_enrichment": ("KEGG_Pathway", "_KEGG", "kegg"),
}

IMAGE_EXTENSIONS = {".png", ".svg", ".jpg", ".jpeg"}
TABLE_EXTENSIONS = {".xlsx", ".xls", ".csv", ".tsv", ".txt"}

METHOD_EXPLANATIONS = {
    "fastqc": "FastQC reports per-base quality, adapter content, duplication, and other FASTQ quality checks.",
    "trim_galore": "Trim Galore trims adapters and low-quality bases before alignment.",
    "trimmomatic": "Trimmomatic trims adapters and low-quality sequence using configurable read filters.",
    "flexbar": "Flexbar performs adapter trimming and read preprocessing.",
    "star": "STAR aligns RNA-seq reads to a genome while supporting splice-aware mapping.",
    "hisat2": "HISAT2 performs fast splice-aware alignment against a graph-based genome index.",
    "bowtie2": "Bowtie2 aligns reads to a reference genome and is useful for general short-read mapping.",
    "bwa": "BWA aligns reads to a reference genome using Burrows-Wheeler indexing.",
    "minimap2": "Minimap2 supports fast alignment for long reads and some short-read modes.",
    "featurecounts": "featureCounts summarizes aligned reads over genomic features.",
    "htseq": "HTSeq counts reads overlapping annotated genomic features.",
    "genomic_overlaps": "PySeqRNA genomic-overlaps quantification counts reads overlapping annotated features in Python.",
    "rpkm": "RPKM normalizes counts by gene length and mapped read depth.",
    "fpkm": "FPKM normalizes paired-end fragment counts by gene length and mapped fragment depth.",
    "tpm": "TPM normalizes transcript abundance so each sample sums to one million transcripts.",
    "cpm": "CPM normalizes by library size as counts per million mapped reads.",
    "median_ratio": "Median-ratio normalization estimates sample-specific size factors from gene count ratios.",
    "tmm": "TMM normalization estimates composition-aware library scaling factors.",
    "deseq2": "DESeq2 performs negative-binomial differential expression analysis through R/Bioconductor.",
    "edger": "edgeR performs negative-binomial differential expression analysis through R/Bioconductor.",
    "pydiffexpress": "PyDiffExpress performs native Python negative-binomial differential expression analysis.",
}


COLUMN_GLOSSARY = {
    "Gene": "Gene or feature identifier.",
    "Name": "Optional gene symbol or short gene name.",
    "Description": "Optional gene, term, or pathway description.",
    "baseMean": "Mean normalized count across all samples.",
    "baseVariance": "Variance of normalized counts across samples.",
    "logFC": "Log2 fold-change for the comparison.",
    "log2FoldChange": "Log2 fold-change for the comparison.",
    "lfcSE": "Standard error of the log2 fold-change estimate.",
    "stat": "Test statistic used to calculate the p-value.",
    "LR": "Likelihood-ratio statistic.",
    "pvalue": "Nominal p-value for the test.",
    "Pvalues": "Nominal p-value for enrichment or differential testing.",
    "padj": "Multiple-testing adjusted p-value.",
    "FDR": "False discovery rate adjusted p-value.",
    "logCPM": "Average log counts per million.",
    "Sample": "Sample identifier.",
    "Total_Reads": "Total reads considered for the sample.",
    "Aligned_Reads": "Reads aligned to the reference genome.",
    "Alignment_Rate": "Percentage of reads aligned to the reference genome.",
    "Mapped_Reads": "Reads mapped to the reference genome or feature set.",
    "Mapped_Rate": "Percentage of reads mapped to the reference genome or feature set.",
    "Uniquely_Mapped_Reads": "Reads mapping uniquely to one genomic location.",
    "Multi_Mapped_Reads": "Reads mapping to multiple genomic locations.",
    "Read_Length": "Read length recorded for the sample or FASTQ file.",
    "Raw_Reads": "Input reads before trimming.",
    "Clean_Reads": "Reads retained after trimming or filtering.",
    "Q20": "Percentage or count of bases with Phred quality score at least 20.",
    "Q30": "Percentage or count of bases with Phred quality score at least 30.",
    "GC": "GC content percentage.",
    "GC_Content": "GC content percentage.",
    "GO_ID": "Gene Ontology identifier.",
    "GO ID": "Gene Ontology identifier.",
    "GO_Term": "Gene Ontology term description.",
    "GO Term": "Gene Ontology term description.",
    "Pathway_ID": "KEGG pathway identifier.",
    "Pathway ID": "KEGG pathway identifier.",
    "Ontology": "Ontology category.",
    "GeneRatio": "Ratio of input genes annotated to the term.",
    "BgRatio": "Ratio of background genes annotated to the term.",
    "Counts": "Number of genes associated with the result row.",
    "Genes": "Gene identifiers contributing to the result row.",
    "logPvalues": "Enrichment score calculated from p-values.",
    "condition": "Experimental condition or treatment group.",
    "replicate": "Biological or technical replicate identifier.",
    "Identifier": "Short condition identifier used for comparisons.",
    "File": "Input or output file path.",
    "File1": "First mate FASTQ file for paired-end samples.",
    "File2": "Second mate FASTQ file for paired-end samples.",
    "SizeFactor": "Sample-specific library size normalization factor.",
    "size_factor": "Sample-specific library size normalization factor.",
    "dispersion": "Negative-binomial dispersion estimate.",
    "AveLogCPM": "Average log counts per million for the gene.",
}


FILE_PATTERNS = [
    ("Raw_Counts", "Raw count matrix from quantification."),
    ("normalized_counts", "Normalized count matrix."),
    ("All_gene_expression", "All-gene differential expression results."),
    ("Filtered_DEGs_summary", "Summary of filtered differentially expressed genes."),
    ("Filtered_DEGs", "Filtered differentially expressed genes."),
    ("Filtered_upDEGs", "Filtered up-regulated genes."),
    ("Filtered_downDEGs", "Filtered down-regulated genes."),
    ("Filtered_DEG", "Differential expression summary plot."),
    ("alignment_stats", "Alignment statistics output."),
    ("alignment_statistics", "Alignment statistics output."),
    ("Aligned.out.bam", "Aligned reads in BAM format."),
    (".bam", "Aligned reads in BAM format."),
    (".bai", "BAM index file."),
    ("trimming_statistics", "Trimming statistics output."),
    ("fastqc", "FastQC quality report or output."),
    ("GO", "Gene Ontology enrichment result or plot."),
    ("KEGG", "KEGG pathway enrichment result or plot."),
    ("Volcano", "Volcano plot."),
    ("MA", "MA plot."),
    ("Heatmap", "Heatmap plot."),
    ("PCA", "Principal component analysis plot."),
    ("tsne", "t-SNE plot."),
    ("cluster", "Clustering output."),
    ("dry_run_report", "Dry-run execution plan generated by PySeqRNA."),
]


@dataclass
class ReportContext:
    output_dir: Path
    report_dir: Path
    title: str = "PySeqRNA Analysis Report"
    input_file: Optional[str] = None
    samples_path: Optional[str] = None
    reference_genome: Optional[str] = None
    feature_file: Optional[str] = None
    config: Dict[str, Any] = field(default_factory=dict)
    checkpoint: Dict[str, Any] = field(default_factory=dict)
    sample_table: Optional[pd.DataFrame] = None
    inventory: List[Dict[str, Any]] = field(default_factory=list)
    previews: Dict[str, Dict[str, Any]] = field(default_factory=dict)


class ReportGenerator:
    """Generate HTML, Markdown, JSON, DOCX, and PDF reports for a PySeqRNA run."""

    def __init__(
        self,
        output_dir: str | Path,
        report_dir: str | Path | None = None,
        title: str = "PySeqRNA Analysis Report",
        input_file: Optional[str] = None,
        samples_path: Optional[str] = None,
        reference_genome: Optional[str] = None,
        feature_file: Optional[str] = None,
        config: Optional[Dict[str, Any]] = None,
        max_preview_rows: int = 20,
        logger: Optional[Any] = None,
    ):
        self.output_dir = Path(output_dir).resolve()
        self.report_dir = Path(report_dir).resolve() if report_dir else self.output_dir / "7.Report"
        self.title = title
        self.input_file = input_file
        self.samples_path = samples_path
        self.reference_genome = reference_genome
        self.feature_file = feature_file
        self.config = config or {}
        self.max_preview_rows = max_preview_rows
        self.logger = logger

    def generate(self, formats: Iterable[str] = ("html", "md", "json")) -> Dict[str, str]:
        """Generate the requested report formats and return output paths."""
        formats = self.parse_formats(formats)
        if not formats:
            raise ReportGenerationError("At least one report format is required.")

        self.report_dir.mkdir(parents=True, exist_ok=True)
        context = self._build_context()

        written: Dict[str, str] = {}
        if "json" in formats:
            written["json"] = self._write_json(context)
        if "md" in formats:
            written["md"] = self._write_markdown(context)
        if "html" in formats:
            written["html"] = self._write_html(context)
        if "docx" in formats:
            written["docx"] = self._write_docx(context)
        if "pdf" in formats:
            written["pdf"] = self._write_pdf(context)

        return written

    @staticmethod
    def parse_formats(formats: Iterable[str] | str) -> List[str]:
        """Normalize and validate report format names."""
        if isinstance(formats, str):
            raw_formats = formats.split(",")
        else:
            raw_formats = []
            for fmt in formats:
                raw_formats.extend(str(fmt).split(","))

        normalized = []
        allowed = {"html", "md", "json", "docx", "pdf"}
        for fmt in raw_formats:
            name = str(fmt).strip().lower()
            if not name:
                continue
            if name not in allowed:
                raise ReportGenerationError(
                    f"Unsupported report format '{name}'. Supported formats: {', '.join(sorted(allowed))}"
                )
            if name not in normalized:
                normalized.append(name)
        return normalized

    def _build_context(self) -> ReportContext:
        checkpoint = self._load_checkpoint()
        sample_table = self._load_sample_table(self.input_file)
        inventory = self._build_inventory()
        previews = self._build_previews(inventory)
        return ReportContext(
            output_dir=self.output_dir,
            report_dir=self.report_dir,
            title=self.title,
            input_file=self.input_file,
            samples_path=self.samples_path,
            reference_genome=self.reference_genome,
            feature_file=self.feature_file,
            config=self.config,
            checkpoint=checkpoint,
            sample_table=sample_table,
            inventory=inventory,
            previews=previews,
        )

    def _load_checkpoint(self) -> Dict[str, Any]:
        checkpoint_file = self.output_dir / "pyseqrna_checkpoint.json"
        if not checkpoint_file.exists():
            return {}
        try:
            with checkpoint_file.open("r", encoding="utf-8") as handle:
                return json.load(handle)
        except Exception as exc:
            self._log("warning", f"Could not read checkpoint file: {exc}")
            return {}

    def _load_sample_table(self, input_file: Optional[str]) -> Optional[pd.DataFrame]:
        if not input_file:
            return None
        path = Path(input_file)
        if not path.exists():
            return None
        try:
            suffix = path.suffix.lower()
            if suffix in {".xlsx", ".xls"}:
                return pd.read_excel(path)
            if suffix == ".csv":
                return pd.read_csv(path, comment="#")
            return pd.read_csv(path, sep=r"\s+", comment="#")
        except Exception as exc:
            self._log("warning", f"Could not read sample table for report: {exc}")
            return None

    def _build_inventory(self) -> List[Dict[str, Any]]:
        files: List[Dict[str, Any]] = []
        excluded = {self.report_dir.resolve()}
        for path in sorted(self.output_dir.rglob("*")):
            if not path.is_file():
                continue
            if any(parent == path.resolve() or parent in path.resolve().parents for parent in excluded):
                continue
            rel = path.relative_to(self.output_dir)
            stat = path.stat()
            files.append(
                {
                    "path": str(path),
                    "relative_path": str(rel),
                    "name": path.name,
                    "extension": path.suffix.lower(),
                    "size_kb": round(stat.st_size / 1024, 2),
                    "stage": self._infer_stage(rel),
                    "description": self._describe_file(path),
                }
            )
        return files

    def _build_previews(self, inventory: List[Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
        previews: Dict[str, Dict[str, Any]] = {}
        preview_candidates = sorted(
            [
                item
                for item in inventory
                if Path(item["path"]).suffix.lower() in TABLE_EXTENSIONS
                and "logs/" not in item["relative_path"].replace("\\", "/").lower()
                and not item["name"].lower().endswith((".err", ".out", ".log"))
            ],
            key=lambda item: (
                0 if any(token.lower() in item["name"].lower() for token, _ in FILE_PATTERNS[:18]) else 1,
                item["stage"],
                item["relative_path"],
            ),
        )
        for item in preview_candidates:
            path = Path(item["path"])
            if len(previews) >= 80:
                break
            preview = self._preview_table(path)
            if preview:
                previews[item["relative_path"]] = preview
        return previews

    def _preview_table(self, path: Path) -> Optional[Dict[str, Any]]:
        try:
            suffix = path.suffix.lower()
            if suffix in {".xlsx", ".xls"}:
                xls = pd.ExcelFile(path)
                sheet = xls.sheet_names[0]
                df = pd.read_excel(path, sheet_name=sheet, nrows=self.max_preview_rows)
                total_rows = pd.read_excel(path, sheet_name=sheet, usecols=[0]).shape[0]
                return {
                    "sheet": sheet,
                    "sheets": xls.sheet_names[:20],
                    "columns": [str(col) for col in df.columns],
                    "rows_previewed": len(df),
                    "estimated_rows": int(total_rows),
                    "html": df.to_html(index=False, classes="data-table", border=0),
                    "markdown": df.to_markdown(index=False),
                    "column_descriptions": self._column_descriptions(df.columns),
                }
            if suffix == ".csv":
                df = pd.read_csv(path, nrows=self.max_preview_rows)
            else:
                try:
                    df = pd.read_csv(path, sep="\t", nrows=self.max_preview_rows)
                except Exception:
                    return None
            return {
                "columns": [str(col) for col in df.columns],
                "rows_previewed": len(df),
                "html": df.to_html(index=False, classes="data-table", border=0),
                "markdown": df.to_markdown(index=False),
                "column_descriptions": self._column_descriptions(df.columns),
            }
        except Exception:
            return None

    def _column_descriptions(self, columns: Iterable[Any]) -> Dict[str, str]:
        descriptions = {}
        for col in columns:
            name = str(col)
            clean = name.split("(")[0].strip()
            descriptions[name] = COLUMN_GLOSSARY.get(
                name,
                COLUMN_GLOSSARY.get(clean, "Output column generated by the selected PySeqRNA stage."),
            )
        return descriptions

    def _infer_stage(self, rel_path: Path) -> str:
        text = str(rel_path).lower()
        top = str(rel_path.parts[0]).lower() if rel_path.parts else ""

        if "gene_ontology" in text or "/go_" in text or "go_" in text:
            return "gene_ontology"
        if "kegg" in text or "pathway" in text:
            return "pathway_enrichment"
        if "multimapped" in text or "raw_mmg" in text or "/mmg" in text:
            return "multimapped_groups"
        if "trim_fastqc" in text or "_val_1_fastqc" in text or "_val_2_fastqc" in text:
            return "quality_trim"
        if (
            "trim_galore" in text
            or "trimmomatic" in text
            or "flexbar" in text
            or "trimming_stats" in text
            or "trimming_statistics" in text
        ):
            return "trimming"
        if "fastqc_results" in text or "_fastqc.html" in text or "_fastqc.zip" in text:
            return "quality"
        if "alignment_stats" in text or "alignment_statistics" in text or "_aligned.bam" in text or ".bai" in text:
            return "alignment"

        mapping = [
            ("alignment", "alignment"),
            ("multimapped", "multimapped_groups"),
            ("normalization", "normalization"),
            ("coexpression", "coexpression"),
            ("clust", "coexpression"),
            ("clustering", "sample_clustering"),
            ("cluster", "sample_clustering"),
            ("differential", "differential"),
            ("visualization", "visualization"),
            ("gene_ontology", "gene_ontology"),
            ("kegg", "pathway_enrichment"),
            ("quantification", "quantification"),
            ("trimming", "trimming"),
            ("quality", "quality"),
        ]
        for token, stage in mapping:
            if token in top or token in text:
                return stage
        return "general"

    def _describe_file(self, path: Path) -> str:
        name = path.name
        for token, description in FILE_PATTERNS:
            if token.lower() in name.lower():
                return description
        if path.suffix.lower() in {".png", ".pdf", ".svg", ".jpg", ".jpeg"}:
            return "Plot or figure generated by PySeqRNA."
        if path.suffix.lower() in {".xlsx", ".csv", ".tsv", ".txt"}:
            return "Tabular result file generated by PySeqRNA."
        return "Pipeline output file."

    def _stage_summary(self, context: ReportContext) -> List[Dict[str, Any]]:
        checkpoint = context.checkpoint
        details = checkpoint.get("stage_metadata", {})
        completed = set(checkpoint.get("completed_stages", []))
        stage_order = list(CheckpointManager.AVAILABLE_STAGES.keys())
        summary = []
        for stage in stage_order:
            metadata = details.get(stage, {})
            files = [item for item in context.inventory if item["stage"] == stage]
            artifact_complete = self._stage_has_completion_evidence(stage, files)
            is_complete = stage in completed or artifact_complete
            health, health_reason = self._stage_health(
                stage,
                is_complete,
                files,
                metadata,
                artifact_complete,
                stage in completed,
            )
            summary.append(
                {
                    "stage": stage,
                    "label": CheckpointManager.AVAILABLE_STAGES.get(stage, stage.replace("_", " ").title()),
                    "status": "complete" if is_complete else "not completed",
                    "health": health,
                    "health_reason": health_reason,
                    "tool": metadata.get("tool")
                    or metadata.get("method")
                    or metadata.get("engine")
                    or metadata.get("quant_method")
                    or self._tool_from_config(stage, context.config)
                    or "not recorded",
                    "completion_time": metadata.get("completion_time", "not recorded"),
                    "output": metadata.get("output_directory")
                    or metadata.get("output_file")
                    or metadata.get("output_files")
                    or "see output inventory",
                    "summary_stats": metadata.get("summary_stats", {}),
                    "metadata": metadata,
                    "narrative": STAGE_NARRATIVES.get(stage, "PySeqRNA analysis stage."),
                    "method_explanation": self._method_explanation(metadata, stage, context.config),
                    "action_items": self._action_items(stage, is_complete, files, metadata),
                    "checkpoint_complete": stage in completed,
                    "artifact_complete": artifact_complete,
                }
            )
        return summary

    def _stage_has_completion_evidence(self, stage: str, files: Sequence[Dict[str, Any]]) -> bool:
        """Infer report-stage completion from concrete outputs on disk."""
        if not files:
            return False
        evidence = STAGE_COMPLETION_EVIDENCE.get(stage, ())
        searchable = " ".join(f"{item.get('relative_path', '')} {item.get('name', '')}".lower() for item in files)
        return any(token.lower() in searchable for token in evidence)

    def _stage_health(
        self,
        stage: str,
        completed: bool,
        files: List[Dict[str, Any]],
        metadata: Dict[str, Any],
        artifact_complete: bool = False,
        checkpoint_complete: bool = False,
    ) -> tuple[str, str]:
        """Return a stage health label and short explanation."""
        if not completed:
            return "skip", "Stage was not completed or was skipped."
        if metadata.get("dry_run"):
            return "info", "Stage was completed in dry-run mode."
        if artifact_complete and not checkpoint_complete:
            return (
                "pass",
                "Report detected expected output files even though the checkpoint did not mark this stage complete.",
            )
        if not files and not metadata.get("output_file") and not metadata.get("output_directory"):
            return "warn", "Checkpoint is complete, but no output files were detected."
        if metadata.get("summary_stats") or files:
            return "pass", "Stage completed and reportable outputs were detected."
        return "info", "Stage completed, but limited metadata was available."

    def _action_items(
        self,
        stage: str,
        completed: bool,
        files: List[Dict[str, Any]],
        metadata: Dict[str, Any],
    ) -> List[str]:
        """Generate concise reviewer notes for each stage."""
        if not completed:
            return ["Confirm whether this stage was intentionally skipped or should be rerun."]

        notes = []
        names = " ".join(item["name"].lower() for item in files)
        rels = " ".join(item["relative_path"].lower() for item in files)
        combined = f"{names} {rels}"

        if stage in {"quality", "quality_trim"} and "fastqc" not in combined:
            notes.append("No FastQC output was detected; verify quality control outputs if this stage was expected.")
        if stage == "trimming" and not any(token in combined for token in ["trim", "trimmomatic", "flexbar"]):
            notes.append("No trimming report was detected; verify trimmed FASTQ files and trimming statistics.")
        if stage == "alignment" and not any(
            token in combined for token in [".bam", "alignment_stats", "alignment_statistics"]
        ):
            notes.append(
                "No BAM or alignment statistics file was detected; inspect aligner output before downstream interpretation."
            )
        if stage == "quantification" and "raw_counts" not in combined:
            notes.append(
                "No raw count matrix was detected; downstream normalization and differential expression may be unavailable."
            )
        if (
            stage == "normalization"
            and "normalized" not in combined
            and not any(token in combined for token in ["rpkm", "tpm", "fpkm", "cpm", "median", "tmm"])
        ):
            notes.append("No normalized count matrix was detected; PCA, clustering, and heatmaps may be limited.")
        if stage == "differential" and "filtered_degs" not in combined:
            notes.append("No filtered DEG table was detected; check thresholds or differential expression output files.")
        if stage == "visualization" and not any(
            item["extension"] in {".png", ".pdf", ".svg", ".jpg", ".jpeg"} for item in files
        ):
            notes.append("No visualization image files were detected for this stage.")
        if stage in {"gene_ontology", "pathway_enrichment"} and not files:
            notes.append("No enrichment outputs were detected; this may be expected if functional annotation was skipped.")

        if not notes:
            notes.append("Review the listed output files and summary statistics for biological interpretation.")
        return notes

    def _dashboard_metrics(self, context: ReportContext, stage_summary: List[Dict[str, Any]]) -> List[List[str]]:
        """Build high-level report metrics similar to a compact general-stats table."""
        total_stages = len(stage_summary)
        completed = sum(1 for stage in stage_summary if stage["status"] == "complete")
        warned = sum(1 for stage in stage_summary if stage["health"] == "warn")
        skipped = sum(1 for stage in stage_summary if stage["health"] == "skip")
        total_size_mb = sum(item.get("size_kb", 0) for item in context.inventory) / 1024
        plot_count = sum(
            1 for item in context.inventory if item["extension"] in IMAGE_EXTENSIONS or item["extension"] == ".pdf"
        )
        table_count = sum(1 for item in context.inventory if item["extension"] in TABLE_EXTENSIONS)
        sample_count = context.sample_table.shape[0] if context.sample_table is not None else "not provided"
        return [
            [
                "Samples",
                str(sample_count),
                "Sample rows detected from the input sheet.",
            ],
            [
                "Stages complete",
                f"{completed}/{total_stages}",
                "Checkpoint plus output-file completion summary.",
            ],
            [
                "Warnings",
                str(warned),
                "Completed stages with limited or missing detected outputs.",
            ],
            [
                "Skipped/incomplete",
                str(skipped),
                "Stages not marked complete in the checkpoint.",
            ],
            [
                "Output files",
                str(len(context.inventory)),
                "Files detected under the run directory.",
            ],
            ["Output size", f"{total_size_mb:.2f} MB", "Total detected output size."],
            ["Tables", str(table_count), "Detected tabular outputs."],
            ["Plots", str(plot_count), "Detected plot/image outputs."],
        ]

    def _tool_from_config(self, stage: str, config: Dict[str, Any]) -> Optional[str]:
        key_map = {
            "quality": "quality_tool",
            "quality_trim": "quality_tool",
            "trimming": "trimming_tool",
            "alignment": "alignment_tool",
            "quantification": "quant_method",
            "normalization": "normalization_method",
            "sample_clustering": "cluster_method",
            "coexpression": "coexpression_tool",
            "differential": "diffexp_tool",
        }
        return config.get(key_map.get(stage, ""))

    def _method_explanation(self, metadata: Dict[str, Any], stage: str, config: Dict[str, Any]) -> str:
        tool = (
            metadata.get("tool")
            or metadata.get("method")
            or metadata.get("engine")
            or metadata.get("quant_method")
            or self._tool_from_config(stage, config)
            or ""
        )
        return METHOD_EXPLANATIONS.get(
            str(tool).lower(),
            "No method-specific explanation was recorded for this stage.",
        )

    def _write_json(self, context: ReportContext) -> str:
        path = context.report_dir / "pyseqrna_report.json"
        payload = {
            "title": context.title,
            "generated_at": datetime.now().isoformat(),
            "pyseqrna_version": __version__,
            "system": self._system_info(),
            "inputs": self._input_summary(context),
            "config": self._json_safe(context.config),
            "checkpoint": self._json_safe(context.checkpoint),
            "stage_summary": self._json_safe(self._stage_summary(context)),
            "dashboard_metrics": self._json_safe(self._dashboard_metrics(context, self._stage_summary(context))),
            "output_inventory": self._json_safe(context.inventory),
            "table_previews": {
                key: {k: v for k, v in value.items() if k not in {"html", "markdown"}}
                for key, value in context.previews.items()
            },
        }
        path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
        return str(path)

    def _write_markdown(self, context: ReportContext) -> str:
        path = context.report_dir / "pyseqrna_report.md"
        stage_summary = self._stage_summary(context)
        lines = [
            f"# {context.title}",
            "",
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"PySeqRNA version: {__version__}",
            "",
            "## Run Dashboard",
            "",
            self._markdown_table(
                self._dashboard_metrics(context, stage_summary),
                ["Metric", "Value", "Notes"],
            ),
            "",
            "## Analysis Inputs",
            "",
            self._markdown_table(self._input_rows(context), ["Item", "Value"]),
            "",
            "## Stage Summary",
            "",
            self._markdown_table(
                [
                    [
                        s["label"],
                        s["status"],
                        s["health"],
                        s["tool"],
                        self._short_value(s["output"]),
                    ]
                    for s in stage_summary
                ],
                ["Stage", "Status", "Health", "Tool/Method", "Main Output"],
            ),
            "",
            "## Stage Details",
            "",
        ]
        for stage in stage_summary:
            lines.extend([f"### {stage['label']}", "", stage["narrative"], ""])
            if stage["summary_stats"]:
                lines.extend(
                    [
                        "Summary statistics:",
                        "",
                        self._dict_as_markdown(stage["summary_stats"]),
                        "",
                    ]
                )
            lines.append(f"Method note: {stage['method_explanation']}")
            lines.append(f"Health: {stage['health']} - {stage['health_reason']}")
            lines.append("")
            lines.extend(["Suggested checks:", ""])
            lines.extend([f"- {item}" for item in stage["action_items"]])
            lines.append("")
            files = [item for item in context.inventory if item["stage"] == stage["stage"]]
            if files:
                lines.extend(
                    [
                        "Key output files:",
                        "",
                        self._markdown_table(
                            [
                                [
                                    f["relative_path"],
                                    f["description"],
                                    f"{f['size_kb']} KB",
                                ]
                                for f in files[:20]
                            ],
                            ["File", "Description", "Size"],
                        ),
                        "",
                    ]
                )

        lines.extend(["## Output Inventory", ""])
        lines.append(
            self._markdown_table(
                [
                    [
                        item["stage"],
                        item["relative_path"],
                        item["description"],
                        f"{item['size_kb']} KB",
                    ]
                    for item in context.inventory
                ],
                ["Stage", "File", "Description", "Size"],
            )
        )
        lines.append("")

        lines.extend(["## Output File Column Guide", ""])
        for rel, preview in context.previews.items():
            lines.extend([f"### {rel}", ""])
            lines.append(
                self._markdown_table(
                    [[col, desc] for col, desc in preview.get("column_descriptions", {}).items()],
                    ["Column", "Description"],
                )
            )
            lines.extend(["", "Preview:", "", preview.get("markdown", ""), ""])

        path.write_text("\n".join(lines), encoding="utf-8")
        return str(path)

    def _write_html(self, context: ReportContext) -> str:
        path = context.report_dir / "pyseqrna_report.html"
        stage_summary = self._stage_summary(context)
        dashboard_cards = self._dashboard_cards_html(self._dashboard_metrics(context, stage_summary))
        nav = "\n".join(f'<a href="#{stage["stage"]}">{html.escape(stage["label"])}</a>' for stage in stage_summary)
        stage_cards = []
        for stage in stage_summary:
            files = [item for item in context.inventory if item["stage"] == stage["stage"]]
            plots = self._stage_plots(files)
            file_rows = "\n".join(
                "<tr>"
                f"<td><a href='{html.escape(self._relative_report_link(context, f['path']))}'>{html.escape(f['relative_path'])}</a></td>"
                f"<td>{html.escape(f['description'])}</td>"
                f"<td>{f['size_kb']} KB</td>"
                "</tr>"
                for f in files[:30]
            )
            stats_html = (
                self._dict_as_html(stage["summary_stats"])
                if stage["summary_stats"]
                else "<p>No summary statistics recorded.</p>"
            )
            action_items = "\n".join(f"<li>{html.escape(item)}</li>" for item in stage["action_items"])
            plot_gallery = self._plot_gallery_html(context, plots[:24])
            stage_cards.append(
                f"""
                <section id="{stage["stage"]}" class="card">
                  <h2>{html.escape(stage["label"])}</h2>
                  <p>
                    <span class="status {stage["status"].replace(" ", "-")}">{html.escape(stage["status"].title())}</span>
                    <span class="health health-{stage["health"]}">{html.escape(stage["health"].upper())}</span>
                  </p>
                  <p>{html.escape(stage["narrative"])}</p>
                  <p><strong>Method note:</strong> {html.escape(stage["method_explanation"])}</p>
                  <p><strong>Health note:</strong> {html.escape(stage["health_reason"])}</p>
                  <p><strong>Tool/method:</strong> {html.escape(str(stage["tool"]))}</p>
                  <p><strong>Completion time:</strong> {html.escape(str(stage["completion_time"]))}</p>
                  <h3>Suggested Checks</h3>
                  <ul>{action_items}</ul>
                  <h3>Summary Statistics</h3>
                  {stats_html}
                  <h3>Plots and Figures</h3>
                  {plot_gallery}
                  <h3>Output Files</h3>
                  <table><thead><tr><th>File</th><th>Description</th><th>Size</th></tr></thead><tbody>{file_rows}</tbody></table>
                </section>
                """
            )

        preview_sections = []
        for rel, preview in context.previews.items():
            col_rows = "\n".join(
                f"<tr><td>{html.escape(col)}</td><td>{html.escape(desc)}</td></tr>"
                for col, desc in preview.get("column_descriptions", {}).items()
            )
            preview_sections.append(
                f"""
                <section class="card">
                  <h3>{html.escape(rel)}</h3>
                  <p><a href="{html.escape(self._relative_report_link(context, context.output_dir / rel))}">Open result file</a></p>
                  <h4>Column Guide</h4>
                  <table><thead><tr><th>Column</th><th>Description</th></tr></thead><tbody>{col_rows}</tbody></table>
                  <h4>Preview</h4>
                  <div class="table-wrap">{preview.get("html", "")}</div>
                </section>
                """
            )

        html_text = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(context.title)}</title>
  <style>
    :root {{ --ink:#18212f; --muted:#607084; --line:#dbe3ee; --brand:#0f766e; --bg:#f6f8fb; --card:#ffffff; }}
    body {{ margin:0; font-family: Arial, sans-serif; color:var(--ink); background:var(--bg); line-height:1.55; }}
    header {{ padding:38px 42px; background:linear-gradient(135deg,#0f766e,#123047); color:white; }}
    header h1 {{ margin:0; font-size:34px; }}
    header p {{ margin:8px 0 0; color:#d9fffb; }}
    .layout {{ display:grid; grid-template-columns:280px 1fr; gap:24px; padding:24px; }}
    nav {{ position:sticky; top:20px; align-self:start; background:var(--card); border:1px solid var(--line); border-radius:16px; padding:18px; }}
    nav a {{ display:block; padding:7px 0; color:var(--brand); text-decoration:none; }}
    main {{ min-width:0; }}
    .card {{ background:var(--card); border:1px solid var(--line); border-radius:18px; padding:22px; margin-bottom:22px; box-shadow:0 12px 28px rgba(24,33,47,0.06); }}
    .metric-grid {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(170px,1fr)); gap:14px; margin-top:18px; }}
    .metric {{ border:1px solid var(--line); border-radius:14px; padding:15px; background:linear-gradient(180deg,#ffffff,#f8fbfb); }}
    .metric .label {{ color:var(--muted); font-size:12px; text-transform:uppercase; letter-spacing:.04em; }}
    .metric .value {{ display:block; font-size:26px; font-weight:800; color:#123047; margin:5px 0; }}
    .metric .note {{ color:var(--muted); font-size:12px; }}
    h2 {{ margin-top:0; }}
    table {{ border-collapse:collapse; width:100%; margin:12px 0; font-size:13px; }}
    th, td {{ border:1px solid var(--line); padding:8px; text-align:left; vertical-align:top; }}
    th {{ background:#eef6f5; }}
    .table-wrap {{ overflow:auto; max-height:480px; border:1px solid var(--line); border-radius:12px; }}
    .status {{ display:inline-block; padding:4px 10px; border-radius:999px; font-weight:bold; font-size:12px; }}
    .health {{ display:inline-block; padding:4px 10px; border-radius:999px; font-weight:bold; font-size:12px; margin-left:6px; }}
    .complete {{ background:#dcfce7; color:#166534; }}
    .not-completed {{ background:#fef3c7; color:#92400e; }}
    .health-pass {{ background:#dcfce7; color:#166534; }}
    .health-warn {{ background:#fee2e2; color:#991b1b; }}
    .health-skip {{ background:#f1f5f9; color:#475569; }}
    .health-info {{ background:#dbeafe; color:#1e40af; }}
    .filter-input {{ width:100%; box-sizing:border-box; border:1px solid var(--line); border-radius:12px; padding:11px 13px; font-size:14px; }}
    .plot-grid {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(260px,1fr)); gap:16px; margin:12px 0; }}
    .plot-card {{ border:1px solid var(--line); border-radius:14px; background:#fbfdff; overflow:hidden; }}
    .plot-card img {{ display:block; width:100%; height:auto; background:white; }}
    .plot-card figcaption {{ padding:9px 11px; font-size:12px; color:var(--muted); border-top:1px solid var(--line); word-break:break-word; }}
    code {{ background:#edf2f7; padding:2px 5px; border-radius:5px; }}
    @media (max-width:900px) {{ .layout {{ grid-template-columns:1fr; }} nav {{ position:relative; top:auto; }} }}
  </style>
</head>
<body>
  <header>
    <h1>{html.escape(context.title)}</h1>
    <p>Generated {datetime.now().strftime("%Y-%m-%d %H:%M:%S")} with PySeqRNA {__version__}</p>
  </header>
  <div class="layout">
    <nav>
      <strong>Report Sections</strong>
      <a href="#inputs">Analysis Inputs</a>
      <a href="#dashboard">Run Dashboard</a>
      <a href="#summary">Stage Summary</a>
      {nav}
      <a href="#inventory">Output Inventory</a>
      <a href="#columns">Output Columns</a>
    </nav>
    <main>
      <section id="dashboard" class="card">
        <h2>Run Dashboard</h2>
        <p>This dashboard gives a quick summary of samples, stage completion, warnings, and detected outputs before the detailed report.</p>
        <div class="metric-grid">{dashboard_cards}</div>
      </section>
      <section id="inputs" class="card">
        <h2>Analysis Inputs</h2>
        {self._rows_as_html_table(self._input_rows(context), ["Item", "Value"])}
        {self._sample_table_html(context.sample_table)}
      </section>
      <section id="summary" class="card">
        <h2>Stage Summary</h2>
        {self._rows_as_html_table([[s["label"], s["status"], s["health"], s["tool"], self._short_value(s["output"])] for s in stage_summary], ["Stage", "Status", "Health", "Tool/Method", "Main Output"])}
      </section>
      {"".join(stage_cards)}
      <section id="inventory" class="card">
        <h2>Output Inventory</h2>
        <p>Filter outputs by stage, filename, or description.</p>
        <input id="inventoryFilter" class="filter-input" type="search" placeholder="Filter output files...">
        {self._inventory_table_html(context)}
      </section>
      <section id="columns" class="card">
        <h2>Output File Column Guide and Previews</h2>
        <p>The following sections explain detected tabular output files and show a small preview of each file.</p>
      </section>
      {"".join(preview_sections)}
    </main>
  </div>
  <script>
    const filter = document.getElementById('inventoryFilter');
    if (filter) {{
      filter.addEventListener('input', () => {{
        const query = filter.value.toLowerCase();
        document.querySelectorAll('#inventoryTable tbody tr').forEach((row) => {{
          row.style.display = row.innerText.toLowerCase().includes(query) ? '' : 'none';
        }});
      }});
    }}
  </script>
</body>
</html>
"""
        path.write_text(html_text, encoding="utf-8")
        return str(path)

    def _write_docx(self, context: ReportContext) -> str:
        try:
            from docx import Document
            from docx.shared import Inches
        except Exception as exc:
            raise ReportGenerationError("DOCX report requires python-docx. Install with: pip install python-docx") from exc

        path = context.report_dir / "pyseqrna_report.docx"
        stage_summary = self._stage_summary(context)
        doc = Document()
        doc.add_heading(context.title, 0)
        doc.add_paragraph(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} with PySeqRNA {__version__}.")

        doc.add_heading("Run Dashboard", level=1)
        self._docx_table(
            doc,
            self._dashboard_metrics(context, stage_summary),
            ["Metric", "Value", "Notes"],
        )

        doc.add_heading("Analysis Inputs", level=1)
        self._docx_table(doc, self._input_rows(context), ["Item", "Value"])
        if context.sample_table is not None:
            doc.add_heading("Sample Information", level=2)
            self._docx_dataframe(doc, context.sample_table.head(self.max_preview_rows))

        doc.add_heading("Stage Summary", level=1)
        self._docx_table(
            doc,
            [
                [
                    s["label"],
                    s["status"],
                    s["health"],
                    s["tool"],
                    self._short_value(s["output"]),
                ]
                for s in stage_summary
            ],
            ["Stage", "Status", "Health", "Tool/Method", "Main Output"],
        )

        for stage in stage_summary:
            doc.add_heading(stage["label"], level=1)
            doc.add_paragraph(stage["narrative"])
            doc.add_paragraph(f"Method note: {stage['method_explanation']}")
            doc.add_paragraph(f"Health: {stage['health']} - {stage['health_reason']}")
            doc.add_paragraph(f"Status: {stage['status']}")
            doc.add_paragraph(f"Tool/method: {stage['tool']}")
            doc.add_heading("Suggested Checks", level=2)
            for item in stage["action_items"]:
                doc.add_paragraph(item, style="List Bullet")
            if stage["summary_stats"]:
                doc.add_heading("Summary Statistics", level=2)
                self._docx_table(
                    doc,
                    [[k, self._short_value(v)] for k, v in stage["summary_stats"].items()],
                    ["Metric", "Value"],
                )
            files = [item for item in context.inventory if item["stage"] == stage["stage"]]
            if files:
                doc.add_heading("Output Files", level=2)
                self._docx_table(
                    doc,
                    [[f["relative_path"], f["description"], f"{f['size_kb']} KB"] for f in files[:15]],
                    ["File", "Description", "Size"],
                )

        doc.add_heading("Output File Column Guide", level=1)
        for rel, preview in list(context.previews.items())[:12]:
            doc.add_heading(rel, level=2)
            self._docx_table(
                doc,
                [[col, desc] for col, desc in preview.get("column_descriptions", {}).items()],
                ["Column", "Description"],
            )

        for section in doc.sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

        doc.save(path)
        return str(path)

    def _write_pdf(self, context: ReportContext) -> str:
        try:
            import matplotlib.pyplot as plt
            from matplotlib.backends.backend_pdf import PdfPages
            import matplotlib.image as mpimg
        except Exception as exc:
            raise ReportGenerationError("PDF report requires matplotlib.") from exc

        path = context.report_dir / "pyseqrna_report.pdf"
        pages = self._pdf_pages(context)
        with PdfPages(path) as pdf:
            for title, body in pages:
                fig = plt.figure(figsize=(8.5, 11))
                fig.patch.set_facecolor("white")
                fig.text(0.08, 0.95, title, fontsize=16, weight="bold", va="top")
                fig.text(
                    0.08,
                    0.90,
                    body,
                    fontsize=9,
                    va="top",
                    family="monospace",
                    wrap=True,
                )
                pdf.savefig(fig, bbox_inches="tight")
                plt.close(fig)
            for plot in self._pdf_plot_pages(context, limit=30):
                try:
                    image = mpimg.imread(plot["path"])
                except Exception:
                    continue
                fig, ax = plt.subplots(figsize=(8.5, 11))
                fig.patch.set_facecolor("white")
                ax.imshow(image)
                ax.set_axis_off()
                fig.text(
                    0.08,
                    0.96,
                    plot["relative_path"],
                    fontsize=11,
                    weight="bold",
                    va="top",
                )
                pdf.savefig(fig, bbox_inches="tight")
                plt.close(fig)
        return str(path)

    def _pdf_plot_pages(self, context: ReportContext, limit: int = 30) -> List[Dict[str, Any]]:
        """Select representative raster plots for the PDF report."""
        selected: List[Dict[str, Any]] = []
        for stage in self._stage_summary(context):
            files = [item for item in context.inventory if item["stage"] == stage["stage"]]
            for plot in self._stage_plots(files)[:3]:
                if Path(plot["path"]).suffix.lower() in {".png", ".jpg", ".jpeg"}:
                    selected.append(plot)
                if len(selected) >= limit:
                    return selected
        return selected

    def _pdf_pages(self, context: ReportContext) -> List[tuple[str, str]]:
        stage_summary = self._stage_summary(context)
        pages = [
            (
                context.title,
                "\n".join(
                    [
                        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                        f"PySeqRNA version: {__version__}",
                        "",
                        "Analysis Inputs:",
                        *[f"- {k}: {v}" for k, v in self._input_rows(context)],
                    ]
                ),
            )
        ]
        pages.append(
            (
                "Run Dashboard",
                "\n".join(
                    f"{label}: {value} - {note}" for label, value, note in self._dashboard_metrics(context, stage_summary)
                ),
            )
        )
        stage_lines = [
            f"{s['label']}: {s['status']} | health={s['health']} | tool={s['tool']} | output={self._short_value(s['output'], 90)}"
            for s in stage_summary
        ]
        pages.append(("Stage Summary", "\n".join(stage_lines)))
        for stage in stage_summary:
            files = [item for item in context.inventory if item["stage"] == stage["stage"]]
            body = [
                stage["narrative"],
                "",
                f"Status: {stage['status']}",
                f"Health: {stage['health']} - {stage['health_reason']}",
                f"Tool/method: {stage['tool']}",
                "",
                "Suggested checks:",
                *[f"- {item}" for item in stage["action_items"]],
            ]
            if stage["summary_stats"]:
                body.extend(["", "Summary statistics:"])
                body.extend([f"- {k}: {self._short_value(v, 100)}" for k, v in stage["summary_stats"].items()])
            if files:
                body.extend(["", "Output files:"])
                body.extend([f"- {f['relative_path']} ({f['size_kb']} KB): {f['description']}" for f in files[:20]])
            pages.append((stage["label"], "\n".join(body)))
        return pages

    def _input_summary(self, context: ReportContext) -> Dict[str, Any]:
        return {key: value for key, value in self._input_rows(context)}

    def _input_rows(self, context: ReportContext) -> List[List[str]]:
        checkpoint = context.checkpoint
        stage_summary = self._stage_summary(context)
        completed_count = sum(1 for stage in stage_summary if stage["status"] == "complete")
        rows = [
            ["Output directory", str(context.output_dir)],
            ["Input sample file", context.input_file or "not provided"],
            ["Samples path", context.samples_path or "not provided"],
            ["Reference genome", context.reference_genome or "not provided"],
            ["Feature annotation", context.feature_file or "not provided"],
            ["Completed stages", f"{completed_count}/{len(stage_summary)}"],
            [
                "Dry run",
                str(
                    checkpoint.get(
                        "dry_run_completed",
                        context.config.get("dryrun", "not recorded"),
                    )
                ),
            ],
            ["Threads", str(context.config.get("threads", "not recorded"))],
            ["Memory", str(context.config.get("memory", "not recorded"))],
            ["Report directory", str(context.report_dir)],
        ]
        for key in [
            "quality_tool",
            "trimming_tool",
            "alignment_tool",
            "quant_method",
            "normalization_method",
            "diffexp_tool",
            "diffexp_normalization",
            "diffexp_abundance",
            "diffexp_dispersion",
            "diffexp_test",
            "fdr_threshold",
            "log2fc_threshold",
            "pvalue_threshold",
            "gene_ontology",
            "kegg_pathway",
            "species",
            "organism_type",
        ]:
            if key in context.config and context.config[key] is not None:
                rows.append([key.replace("_", " ").title(), str(context.config[key])])
        if context.sample_table is not None:
            rows.append(["Samples detected", str(context.sample_table.shape[0])])
        return rows

    def _system_info(self) -> Dict[str, str]:
        return {
            "os": f"{platform.system()} {platform.release()}",
            "python": platform.python_version(),
        }

    def _sample_table_html(self, df: Optional[pd.DataFrame]) -> str:
        if df is None:
            return "<p>No sample table was provided to the report generator.</p>"
        return "<h3>Sample Information</h3>" + df.head(self.max_preview_rows).to_html(
            index=False, classes="data-table", border=0
        )

    def _stage_plots(self, files: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Return embeddable plot files, preferring raster images over PDF companions."""
        plots = [item for item in files if item.get("extension") in IMAGE_EXTENSIONS]
        return sorted(plots, key=lambda item: item["relative_path"])

    def _plot_gallery_html(self, context: ReportContext, plots: Sequence[Dict[str, Any]]) -> str:
        if not plots:
            return "<p>No embeddable plot images detected for this stage.</p>"
        cards = []
        for plot in plots:
            link = html.escape(self._relative_report_link(context, plot["path"]))
            caption = html.escape(plot["relative_path"])
            cards.append(
                '<figure class="plot-card">'
                f'<a href="{link}"><img src="{link}" alt="{caption}"></a>'
                f"<figcaption>{caption}</figcaption>"
                "</figure>"
            )
        return '<div class="plot-grid">' + "\n".join(cards) + "</div>"

    def _dashboard_cards_html(self, metrics: List[List[str]]) -> str:
        cards = []
        for label, value, note in metrics:
            cards.append(
                '<div class="metric">'
                f'<span class="label">{html.escape(label)}</span>'
                f'<span class="value">{html.escape(value)}</span>'
                f'<span class="note">{html.escape(note)}</span>'
                "</div>"
            )
        return "\n".join(cards)

    def _inventory_table_html(self, context: ReportContext) -> str:
        rows = []
        for item in context.inventory:
            rows.append(
                "<tr>"
                f"<td>{html.escape(item['stage'])}</td>"
                f"<td><a href='{html.escape(self._relative_report_link(context, item['path']))}'>{html.escape(item['relative_path'])}</a></td>"
                f"<td>{html.escape(item['description'])}</td>"
                f"<td>{html.escape(item['extension'] or 'none')}</td>"
                f"<td>{item['size_kb']} KB</td>"
                "</tr>"
            )
        return (
            '<div class="table-wrap">'
            '<table id="inventoryTable"><thead><tr>'
            "<th>Stage</th><th>File</th><th>Description</th><th>Type</th><th>Size</th>"
            "</tr></thead><tbody>" + "\n".join(rows) + "</tbody></table></div>"
        )

    def _rows_as_html_table(self, rows: List[List[Any]], headers: List[str]) -> str:
        head = "".join(f"<th>{html.escape(str(h))}</th>" for h in headers)
        body = "\n".join(
            "<tr>" + "".join(f"<td>{html.escape(self._short_value(cell))}</td>" for cell in row) + "</tr>" for row in rows
        )
        return f"<table><thead><tr>{head}</tr></thead><tbody>{body}</tbody></table>"

    def _relative_report_link(self, context: ReportContext, target: str | Path) -> str:
        """Return a browser-friendly link from the report directory to a target file."""
        try:
            return Path(target).resolve().relative_to(context.report_dir.resolve()).as_posix()
        except ValueError:
            return (
                Path(target).resolve().relative_to(context.output_dir.resolve()).as_posix()
                if context.report_dir.resolve() == context.output_dir.resolve()
                else Path(os.path.relpath(Path(target).resolve(), context.report_dir.resolve())).as_posix()
            )

    def _dict_as_html(self, data: Dict[str, Any]) -> str:
        return self._rows_as_html_table([[k, self._short_value(v)] for k, v in data.items()], ["Metric", "Value"])

    def _dict_as_markdown(self, data: Dict[str, Any]) -> str:
        return self._markdown_table([[k, self._short_value(v)] for k, v in data.items()], ["Metric", "Value"])

    def _markdown_table(self, rows: List[List[Any]], headers: List[str]) -> str:
        clean_rows = [[self._short_value(cell).replace("\n", " ") for cell in row] for row in rows]
        header = "| " + " | ".join(headers) + " |"
        sep = "| " + " | ".join(["---"] * len(headers)) + " |"
        body = ["| " + " | ".join(row) + " |" for row in clean_rows]
        return "\n".join([header, sep, *body])

    def _docx_table(self, doc: Any, rows: List[List[Any]], headers: List[str]) -> None:
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = str(header)
        for row in rows:
            cells = table.add_row().cells
            for idx, value in enumerate(row[: len(headers)]):
                cells[idx].text = self._short_value(value, 300)

    def _docx_dataframe(self, doc: Any, df: pd.DataFrame) -> None:
        rows = df.fillna("").astype(str).values.tolist()
        self._docx_table(doc, rows, [str(col) for col in df.columns])

    def _short_value(self, value: Any, limit: int = 180) -> str:
        if isinstance(value, (dict, list, tuple)):
            value = json.dumps(self._json_safe(value), ensure_ascii=True)
        text = str(value)
        if len(text) > limit:
            return text[: limit - 3] + "..."
        return text

    def _json_safe(self, value: Any) -> Any:
        if isinstance(value, dict):
            return {str(k): self._json_safe(v) for k, v in value.items()}
        if isinstance(value, list):
            return [self._json_safe(v) for v in value]
        if isinstance(value, tuple):
            return [self._json_safe(v) for v in value]
        if isinstance(value, Path):
            return str(value)
        if hasattr(value, "item"):
            try:
                return value.item()
            except Exception:
                return str(value)
        return value

    def _log(self, level: str, message: str) -> None:
        if self.logger is not None and hasattr(self.logger, level):
            getattr(self.logger, level)(message)


def generate_report(
    output_dir: str | Path,
    formats: Iterable[str] = ("html", "md", "json"),
    report_dir: str | Path | None = None,
    **kwargs: Any,
) -> Dict[str, str]:
    """Convenience wrapper for report generation."""
    return ReportGenerator(output_dir=output_dir, report_dir=report_dir, **kwargs).generate(formats=formats)
